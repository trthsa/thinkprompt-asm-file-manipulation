import os
import pymupdf
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
 
class TreeNode:
    def __init__(self, type, content=None, attributes=None):
        self.type = type  # 'paragraph', 'run', or 'image'
        self.content = content  # Text for 'paragraph' and 'run', path for 'image'
        self.attributes = attributes if attributes else {}
        self.children = []

    def add_child(self, node):
        self.children.append(node)

# Adjusted build_docx_tree function
def build_docx_tree(doc):
    root = TreeNode("document")
    img_dir = "./img"
    os.makedirs(img_dir, exist_ok=True)

    image_count = 0

    for para in doc.paragraphs:
        para_node = TreeNode("paragraph")
        for run in para.runs:
            run_attributes = {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
                "text_color": run.font.color.rgb if run.font.color and run.font.color.rgb else None,
            }
            run_node = TreeNode("run", content=run.text, attributes=run_attributes)
            para_node.add_child(run_node)
        root.add_child(para_node)

    # Note: Due to limitations in identifying exact image positions within paragraphs using python-docx,
    # images are processed separately and added at the end. This may not reflect their original positions.
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_path = os.path.join(img_dir, f"image{image_count}.png")
            image_bytes = rel._target.blob
                 # Save the image data to the specified path
            with open(image_path, 'wb') as img_file:
                img_file.write(image_bytes)
            
            image_node = TreeNode("image", content=image_path)
            # Adjusted to add images to a designated 'images' node for clarity
            images_node = next((child for child in root.children if child.type == "images"), None)
            if not images_node:
                images_node = TreeNode("images")
                root.add_child(images_node)
            images_node.add_child(image_node)
            image_count += 1

    return root

def display_tree(node, level=0):
    indent = " " * (level * 4)
    print(f"{indent}{node.type.upper()}: {node.content if node.content else ''}")
    for attr, value in node.attributes.items():
        print(f"{indent}  {attr}: {value}")
    for child in node.children:
        display_tree(child, level + 1)

def build_pdf_tree(file_path):
    doc = pymupdf.open(file_path)
    pdfByptes = doc.convert_to_pdf()
    pdf= pymupdf.open("pdf",pdfByptes)
    root = TreeNode("document")
    img_dir = "./img"
    os.makedirs(img_dir, exist_ok=True)

    for page_num, page in enumerate(pdf):
        page_node = TreeNode("page", attributes={"number": page_num})
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"]
                        font_name = span["font"]
                        font_size = span["size"]
                        color_int = span["color"]
                        r, g, b = (color_int >> 16) & 0xff, (color_int >> 8) & 0xff, color_int & 0xff
                        color_hex = f"#{r:02x}{g:02x}{b:02x}"
                        text_node = TreeNode("text", content=text, attributes={
                            "font_name": font_name, "font_size": font_size, "color": color_hex
                        })
                        page_node.add_child(text_node)

        image_list = page.get_images(full=True)
        for image_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf.extract_image(xref)
            image_path = os.path.join(img_dir, f"page_{page_num}_image_{image_index}.png")
            with open(image_path, "wb") as img_file:
                img_file.write(base_image["image"])
            image_node = TreeNode("image", content=image_path)
            page_node.add_child(image_node)

        root.add_child(page_node)

    return root


# doc = Document("plain.docx")
# doc_tree = build_docx_tree(doc)
# display_tree(doc_tree)
 

def recreate_docx_from_tree(node, doc=None, current_paragraph=None):
    if node.type == "document":
        doc = Document()
    elif node.type == "paragraph":
        current_paragraph = doc.add_paragraph()
    elif node.type == "run" and current_paragraph is not None:
        run = current_paragraph.add_run(node.content)
        # Apply attributes
        if node.attributes.get("bold"):
            run.bold = True
        if node.attributes.get("italic"):
            run.italic = True
        if node.attributes.get("underline"):
            run.underline = WD_UNDERLINE.SINGLE
        if node.attributes.get("font_name"):
            run.font.name = node.attributes["font_name"]
        if node.attributes.get("font_size"):
            run.font.size = Pt(node.attributes["font_size"])
        if node.attributes.get("text_color"):
            run.font.color.rgb = node.attributes["text_color"]
    elif node.type == "image":
        doc.add_picture(node.content)

    for child in node.children:
        recreate_docx_from_tree(child, doc, current_paragraph)

    return doc

# # # Assuming `doc_tree` is the tree structure created from the original document
# new_doc = recreate_docx_from_tree(doc_tree)
# new_doc.save("plain_recreated.docx")